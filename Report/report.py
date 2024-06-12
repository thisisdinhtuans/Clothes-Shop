from flask import Flask, render_template, send_file
import os
import pandas as pd
from sqlalchemy import create_engine
from docx import Document

app = Flask(__name__)

# Tạo đối tượng kết nối SQLAlchemy
engine = create_engine('postgresql+psycopg2://postgres:Tuan123%406@localhost:5432/store1')
# Hàm tạo báo cáo
def generate_report():
    query = """
        SELECT EXTRACT(YEAR FROM "OrderDate") AS "Year",
            EXTRACT(MONTH FROM "OrderDate") AS "Month",
            "ShippingAddress_State",
            SUM("Subtotal") AS "TotalSales"
        FROM "Orders"
        GROUP BY EXTRACT(YEAR FROM "OrderDate"), EXTRACT(MONTH FROM "OrderDate"), "ShippingAddress_State"
    """
    
    df = pd.read_sql_query(query, engine)
    df['Year'] = df['Year'].astype(int)
    df['Month'] = df['Month'].astype(int)
    df.columns = ['Năm', 'Tháng', 'Tỉnh/Thành phố', 'Doanh số']
    
    template_path = r'D:\PTHTTTQ\template.docx'
    
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"The template file was not found at {template_path}")
    
    template = Document(template_path)
    
    def add_table(paragraph, dataframe):
        table = template.add_table(rows=len(dataframe) + 1, cols=len(dataframe.columns))
        for j, column in enumerate(dataframe.columns):
            table.cell(0, j).text = str(column)
        for i, row in dataframe.iterrows():
            for j, value in enumerate(row):
                if dataframe.columns[j] == 'Tháng':
                    table.cell(i + 1, j).text = f"{int(value):02d}"
                elif dataframe.columns[j] == 'Năm':
                    table.cell(i + 1, j).text = str(int(value))
                elif dataframe.columns[j] == 'Doanh số':
                    table.cell(i + 1, j).text = f"{value:,.0f}"
                else:
                    table.cell(i + 1, j).text = str(value)
        move_table_to_paragraph(table, paragraph)
    
    def move_table_to_paragraph(table, paragraph):
        tbl, p = table._tbl, paragraph._element
        p.addnext(tbl)
        paragraph.text = paragraph.text.replace('{TABLE_YEARLY_SALES}', '')
        paragraph.text = paragraph.text.replace('{TABLE_MONTHLY_SALES}', '')
        paragraph.text = paragraph.text.replace('{TABLE_STATE_SALES}', '')
    
    sales_by_month = df.groupby(['Năm', 'Tháng'])['Doanh số'].sum().reset_index()
    months_achieved_kpi = sales_by_month[sales_by_month['Doanh số'] > 50000000]
    months_not_achieved_kpi = sales_by_month[sales_by_month['Doanh số'] <= 50000000]
    
    months_achieved_kpi_text = ""
    if not months_achieved_kpi.empty:
        months_achieved_kpi_text += "Những tháng đạt KPI: " + ", ".join(f"{int(row['Tháng']):02d}/{int(row['Năm'])}" for _, row in months_achieved_kpi.iterrows())
    
    months_not_achieved_kpi_text = ""
    if not months_not_achieved_kpi.empty:
        months_not_achieved_kpi_text += "Những tháng không đạt KPI: " + ", ".join(f"{int(row['Tháng']):02d}/{int(row['Năm'])}" for _, row in months_not_achieved_kpi.iterrows())
    
    for paragraph in template.paragraphs:
        if '{TABLE_YEARLY_SALES}' in paragraph.text:
            sales_by_year = df.groupby('Năm')['Doanh số'].sum().reset_index()
            add_table(paragraph, sales_by_year)
    
        if '{TABLE_MONTHLY_SALES}' in paragraph.text:
            add_table(paragraph, sales_by_month)
    
        if '{TABLE_STATE_SALES}' in paragraph.text:
            sales_by_state = df.groupby('Tỉnh/Thành phố')['Doanh số'].sum().reset_index()
            add_table(paragraph, sales_by_state)
    
        if '{MONTHS_ACHIEVED_KPI}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{MONTHS_ACHIEVED_KPI}', months_achieved_kpi_text)
    
        if '{MONTHS_NOT_ACHIEVED_KPI}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{MONTHS_NOT_ACHIEVED_KPI}', months_not_achieved_kpi_text)
    
    report_path = 'Sales_report.docx'
    template.save(report_path)
    
    return report_path

# Route để tạo báo cáo
@app.route('/generate_report')
def generate_report_route():
    report_path = generate_report()
    return send_file(report_path, as_attachment=True)

if __name__ == '__main__':
    # app.run(debug=True)
    app.run(debug=True, host='127.0.0.1', port=5001)
